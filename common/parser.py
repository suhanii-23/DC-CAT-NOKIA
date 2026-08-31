"""Document parsing: PDF (PyMuPDF) and DOCX (python-docx) -> common.contracts.Document.

Depends on common.contracts only.
"""
from __future__ import annotations

import os
import re
import statistics
from typing import NamedTuple, Optional

import pymupdf as fitz  # PyMuPDF

from common.contracts import Document, Heading, LinkAnnotation, Page, Paragraph

_HEADING_NUMBER_RE = re.compile(r"^\s*((?:\d+\.)*\d+)\.?\s+(.*)$")


def parse(path: str) -> Document:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext == ".docx":
        return _parse_docx(path)
    raise ValueError(f"Unsupported document format: {ext!r}")


def _split_number_title(text: str) -> tuple[Optional[str], str]:
    match = _HEADING_NUMBER_RE.match(text)
    if match:
        return match.group(1), match.group(2).strip()
    return None, text.strip()


# --- paragraph splitting ----------------------------------------------
#
# page.get_text("text") returns one line per *rendered* line and, in every
# layout measured against this repo's fixtures, no blank lines at all — so
# splitting a page on "\n\n" left the whole page as one paragraph and the
# keyword-search features searched page-sized blocks.
#
# PyMuPDF's own block grouping is not enough on its own. Measured over
# generated PDFs at a range of font sizes and leadings:
#
#   paragraphs separated by extra space  -> one block per paragraph (good)
#   uniform leading, no space between    -> ONE BLOCK PER PAGE (the bug)
#   leading above about 1.6x             -> one block per LINE (shreds
#                                           sentences across paragraphs)
#
# So the splitter works from the geometry PyMuPDF exposes per line in
# get_text("dict") — the same lines get_text("text") returns, plus their
# bounding boxes — and applies five rules. A paragraph ends when:
#
#   0. the extracted text has a blank line here. The original "\n\n" rule,
#      kept explicit so it still wins outright when a PDF does supply one;
#   1. the step down to the next line is larger than the page's own line
#      pitch, or moves up/sideways (a new column). This is the "extra
#      space between paragraphs" signal, measured directly rather than
#      taken from PyMuPDF's block grouping, which the table above shows
#      is unreliable in both directions;
#   2. the line ends short of the text column *and* closes a sentence —
#      the ordinary last line of a paragraph;
#   3. the next line opens a list item or a numbered heading;
#   4. the line is itself a numbered heading that ended short.
#
# Rules 2-4 are what recover paragraphs on a page with no vertical
# spacing to read, which is exactly what rules 0 and 1 cannot see.
#
# Deliberately NOT used: capitalisation of the following line. It looks
# attractive and scores well on tidy prose, but it tears real paragraphs
# apart wherever a wrapped line happens to be followed by a proper noun
# ("...restart and the / Discovery Center Console then refreshes..."), and
# a torn paragraph breaks phrase matching for both keyword-search
# features. Merging two paragraphs only makes a search unit coarser;
# tearing one loses matches, so every rule here errs towards merging.
#
# Everything is read off the extracted page, so the same PDF always
# yields the same paragraphs.

# A line whose right edge falls at least this fraction of the column short
# ended of its own accord rather than being wrapped. Anything from 0.02 to
# 0.10 behaves the same on the corpus this was measured against; 0.10 has
# the fewest torn paragraphs.
_RAGGED_LINE_FRACTION = 0.10

# A vertical step larger than this multiple of the page's own line pitch
# is space between paragraphs rather than ordinary leading. Measured
# behaviour is flat from 1.15 to 2.0.
_PARAGRAPH_GAP_FACTOR = 1.35

# A line ending like this closed a sentence.
_SENTENCE_END_RE = re.compile(r"[.!?:;][\"'’”)\]]?$")

# Bullets and list labels: "- ", "* ", "1. ", "2) ", "(a) ".
_LIST_START_RE = re.compile(
    r"^(?:[-–—*•●·]|\(?[0-9]{1,3}[.)]|\(?[a-zA-Z][.)])\s"
)


class _Line(NamedTuple):
    """One rendered line: its text, and the geometry needed to place it."""

    text: str
    left: float
    right: float
    top: float


def _page_lines(page: "fitz.Page") -> list[_Line]:
    """Rendered lines with their geometry, in reading order.

    get_text("dict") exposes the same lines as get_text("text") together
    with their bounding boxes. Returns [] when the extractor cannot supply
    them, which sends the caller to the blank-line fallback.
    """
    try:
        data = page.get_text("dict")
    except Exception:
        return []

    lines: list[_Line] = []
    for block in data.get("blocks", []):
        if block.get("type") != 0:
            continue  # image block: carries no text
        for line in block.get("lines", []):
            text = "".join(span.get("text", "") for span in line.get("spans", []))
            text = text.strip()
            bbox = line.get("bbox")
            if not text or not bbox:
                continue
            left, top, right, _bottom = bbox
            lines.append(_Line(text, left, right, top))
    return lines


def _line_pitch(lines: list[_Line]) -> float:
    """The page's ordinary line-to-line step, ignoring column jumps."""
    steps = [
        later.top - earlier.top
        for earlier, later in zip(lines, lines[1:])
        if later.top > earlier.top
    ]
    return statistics.median(steps) if steps else 0.0


def _blank_line_breaks(text: str, line_count: int) -> set[int]:
    """Positions after which the extracted text carried a blank line.

    The rendered lines and the non-empty lines of get_text("text") are the
    same sequence in the same order, so a blank line in the text maps onto
    a position in that sequence. Keeping this explicit preserves the
    original "\\n\\n" rule: a blank line always ends a paragraph, whatever
    the geometry says. Returns an empty set if the two disagree, in which
    case the geometry is trusted on its own.
    """
    breaks: set[int] = set()
    position = -1
    blank_seen = False
    for raw_line in text.splitlines():
        if raw_line.strip():
            if blank_seen and position >= 0:
                breaks.add(position)
            position += 1
            blank_seen = False
        else:
            blank_seen = True
    if position + 1 != line_count:
        return set()
    return breaks


def _is_numbered_heading(line: str) -> bool:
    """A section number followed by a title, e.g. "3.1 Access Control".

    A wrapped list item is not excluded here because it does not need to
    be: the caller only breaks after a heading whose line ended short, and
    a list item that continues onto the next line runs to the full column.
    """
    if _SENTENCE_END_RE.search(line):
        return False
    return bool(_HEADING_NUMBER_RE.match(line))


def _opens_paragraph(line: str) -> bool:
    """Whether a line starts a paragraph regardless of what precedes it."""
    return bool(_LIST_START_RE.match(line)) or _is_numbered_heading(line)


def _page_paragraphs(page: "fitz.Page", text: str) -> list[str]:
    """Paragraph texts for one page, in reading order."""
    lines = _page_lines(page)
    if not lines:
        # No usable geometry: fall back to the original blank-line split so
        # a page is never dropped.
        return [block.strip() for block in text.split("\n\n") if block.strip()]

    column_right = max(line.right for line in lines)
    column_width = column_right - min(line.left for line in lines) or 1.0
    pitch = _line_pitch(lines)
    blank_breaks = _blank_line_breaks(text, len(lines))

    paragraphs: list[str] = []
    current: list[str] = []

    for position, line in enumerate(lines):
        current.append(line.text)
        following = lines[position + 1] if position + 1 < len(lines) else None
        if following is None:
            break

        ragged = line.right < column_right - _RAGGED_LINE_FRACTION * column_width
        step = following.top - line.top
        if (
            position in blank_breaks
            or (pitch > 0 and (step > pitch * _PARAGRAPH_GAP_FACTOR or step <= 0))
            or (ragged and _SENTENCE_END_RE.search(line.text))
            or _opens_paragraph(following.text)
            or (ragged and _is_numbered_heading(line.text))
        ):
            paragraphs.append("\n".join(current))
            current = []

    if current:
        paragraphs.append("\n".join(current))
    return paragraphs


def _parse_pdf(path: str) -> Document:
    doc = fitz.open(path)
    try:
        pages: list[Page] = []
        paragraphs: list[Paragraph] = []
        para_index = 0

        for page_index in range(doc.page_count):
            page = doc[page_index]
            page_number = page_index + 1
            text = page.get_text("text")
            pages.append(Page(number=page_number, text=text))
            for paragraph_text in _page_paragraphs(page, text):
                paragraphs.append(
                    Paragraph(text=paragraph_text, page=page_number, index=para_index)
                )
                para_index += 1

        headings: list[Heading] = []
        for level, title, page_number in doc.get_toc(simple=True):
            number, clean_title = _split_number_title(title)
            headings.append(
                Heading(text=clean_title, level=level, number=number, page=page_number)
            )

        # Named destinations actually defined in this PDF, resolved once for the
        # whole document. Without this every named link looks unresolvable, and
        # broken-link detection flags every cross-reference in a real document:
        # a Nokia user guide with 76 working links reported 76 broken ones.
        try:
            named_destinations = doc.resolve_names() or {}
        except Exception:  # older PyMuPDF, or a malformed name tree
            named_destinations = {}

        links: list[LinkAnnotation] = []
        for page_index in range(doc.page_count):
            page = doc[page_index]
            page_number = page_index + 1
            for link in page.get_links():
                kind = link.get("kind")
                link_text = _extract_link_text(page, link.get("from"))
                if kind == fitz.LINK_GOTO:
                    target_page_index = link.get("page")
                    if target_page_index is None:
                        target_page_index = -1
                    links.append(
                        LinkAnnotation(
                            page=page_number,
                            text=link_text,
                            target_page=target_page_index + 1,
                            target_name=None,
                            is_internal=True,
                        )
                    )
                elif kind == fitz.LINK_NAMED:
                    name = link.get("nameddest") or link.get("name")
                    destination = named_destinations.get(name)
                    resolved_page = None
                    if isinstance(destination, dict):
                        page_index_value = destination.get("page")
                        if isinstance(page_index_value, int) and page_index_value >= 0:
                            resolved_page = page_index_value + 1
                    links.append(
                        LinkAnnotation(
                            page=page_number,
                            text=link_text,
                            # Resolved to a real page when the destination exists
                            # in this document; left as None when it does not,
                            # which is what makes the link broken.
                            target_page=resolved_page,
                            target_name=name,
                            is_internal=True,
                        )
                    )
                # External links (LINK_URI, LINK_LAUNCH, ...) are out of
                # scope for internal broken-link detection and are skipped.

        return Document(
            path=path,
            format="pdf",
            pages=pages,
            paragraphs=paragraphs,
            headings=headings,
            links=links,
            page_count=doc.page_count,
        )
    finally:
        doc.close()


def _extract_link_text(page: "fitz.Page", rect) -> str:
    if rect is None:
        return ""
    try:
        return page.get_textbox(rect).strip()
    except Exception:
        return ""


def _parse_docx(path: str) -> Document:
    from docx import Document as DocxDocument

    docx_doc = DocxDocument(path)
    paragraphs: list[Paragraph] = []
    headings: list[Heading] = []
    full_text_parts: list[str] = []

    for index, para in enumerate(docx_doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(Paragraph(text=text, page=1, index=index))
        full_text_parts.append(text)
        style_name = (para.style.name if para.style is not None else "") or ""
        if style_name.lower().startswith("heading"):
            level_digits = "".join(ch for ch in style_name if ch.isdigit())
            level = int(level_digits) if level_digits else 1
            number, clean_title = _split_number_title(text)
            headings.append(Heading(text=clean_title, level=level, number=number, page=1))

    page = Page(number=1, text="\n".join(full_text_parts))
    return Document(
        path=path,
        format="docx",
        pages=[page],
        paragraphs=paragraphs,
        headings=headings,
        links=[],
        page_count=1,
    )
